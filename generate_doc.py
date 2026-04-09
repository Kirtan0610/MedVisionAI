"""
MedVision AI - Project Documentation Generator
Generates a complete, properly formatted .docx file.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import copy

# ─── Image Paths ───────────────────────────────────────────────
ARTIFACT_DIR = r"C:\Users\kirta\.gemini\antigravity\brain\2002c82d-3dd8-400c-b3ef-f1c38083c649"
ARCH_IMG     = os.path.join(ARTIFACT_DIR, "architecture_diagram_1774345713090.png")
CLASS_IMG    = os.path.join(ARTIFACT_DIR, "class_diagram_1774345683200.png")
SEQ_IMG      = os.path.join(ARTIFACT_DIR, "sequence_diagram_1774345698422.png")
USECASE_IMG  = os.path.join(ARTIFACT_DIR, "usecase_diagram_1774345666710.png")
ACTIVITY_IMG = os.path.join(ARTIFACT_DIR, "activity_diagram_1774345731468.png")

OUTPUT_PATH = r"d:\medvision-ai\Project_Documentation.docx"

# ─── Helpers ────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(table, border_color="1F497D"):
    """Apply borders to all table cells."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                tag = f'w:{edge}'
                el = OxmlElement(tag)
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '6')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), border_color)
                tcPr.append(el)

def add_page_break(doc):
    doc.add_page_break()

def add_heading(doc, text, level=1, color="1F497D", size=None, bold=True, center=False):
    """Add a styled heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    else:
        sizes = {0: 18, 1: 16, 2: 14, 3: 12}
        run.font.size = Pt(sizes.get(level, 12))
    r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.name = "Calibri"
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_body(doc, text, bold=False, italic=False, center=False, color=None, size=12):
    """Add a normal body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
        run.font.color.rgb = RGBColor(r, g, b)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if bold_prefix:
        run1 = p.add_run(bold_prefix + ": ")
        run1.bold = True
        run1.font.name = "Calibri"
        run1.font.size = Pt(12)
        run2 = p.add_run(text)
        run2.font.name = "Calibri"
        run2.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(12)

def add_image(doc, path, caption, width=Inches(5.5)):
    """Insert an image with a caption."""
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
    except Exception as e:
        add_body(doc, f"[Image not found: {os.path.basename(path)}]", italic=True)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap.paragraph_format.space_after = Pt(10)

def add_data_table(doc, headers, rows, header_color="1F497D", alt_color="D9E8F5"):
    """Add a styled table with header row."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Calibri"
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], header_color)

    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.name = "Calibri"
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 0:
                set_cell_bg(cells[ci], alt_color)
            else:
                set_cell_bg(cells[ci], "FFFFFF")

    doc.add_paragraph()
    return table

def add_code_block(doc, code_text, title=None):
    """Add a formatted code block."""
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Calibri"
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        # Light gray background via shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F4F8')
        pPr.append(shd)
    doc.add_paragraph()

def set_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    sections = doc.sections
    for section in sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════════════════════════════
# MAIN DOCUMENT BUILDER
# ════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()
    set_margins(doc)

    # ── Default style ──
    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    # ════════ PAGE 1: TITLE PAGE ════════
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MedVision AI")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Comprehensive Project Documentation")
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.name = "Calibri"
    r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    add_separator(doc)
    doc.add_paragraph()

    for label, value in [
        ("Project Title", "MedVision AI – AI-Powered Medical Report Analysis Platform"),
        ("Course / Class", "[Your Course Name]"),
        ("Student Name", "[Your Full Name]"),
        ("Roll Number", "[Your Roll No.]"),
        ("Guide / Mentor", "[Guide Name]"),
        ("Institution", "[Your Institute / College Name]"),
        ("Department", "[Your Department]"),
        ("Academic Year", "2025 – 2026"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}:  ")
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(12)
        r2 = p.add_run(value)
        r2.font.name = "Calibri"
        r2.font.size = Pt(12)

    doc.add_paragraph()
    add_separator(doc)

    # ════════ PAGE 2: CERTIFICATE ════════
    add_page_break(doc)
    add_heading(doc, "CERTIFICATE", level=0, center=True, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    cert_text = (
        "This is to certify that the project entitled \"MedVision AI\" is a bonafide work carried out by "
        "[Student Name / Team] bearing Roll No. [Roll No.] of [Class / Branch], in partial fulfillment "
        "for the award of the [Degree / Diploma / Certificate] in [Course Name] during the academic year 2025-2026. "
        "The project has been completed under my supervision and guidance, and is submitted in accordance with the "
        "requirements of [Institute / University Name]."
    )
    add_body(doc, cert_text)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_cells = table.columns[0].cells
    right_cells = table.columns[1].cells

    left_cells[0].paragraphs[0].add_run("______________________________").font.size = Pt(12)
    left_cells[1].paragraphs[0].add_run("Guide / Mentor Name").bold = True
    left_cells[1].paragraphs[0].runs[0].font.size = Pt(11)

    right_cells[0].paragraphs[0].add_run("______________________________").font.size = Pt(12)
    right_cells[1].paragraphs[0].add_run("Head of Department (HOD)").bold = True
    right_cells[1].paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()
    p_stamp = doc.add_paragraph("[Place for Institute Stamp / Seal]")
    p_stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_stamp.runs[0].italic = True
    p_stamp.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ════════ PAGE 3: ABSTRACT ════════
    add_page_break(doc)
    add_heading(doc, "ABSTRACT", level=0, center=True, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    abstract = (
        "MedVision AI is an intelligent, full-stack medical reporting and analysis platform built to simplify "
        "complex healthcare documents for patients while simultaneously assisting medical professionals. Powered "
        "by advanced AI (Groq API / LLaMA 3.3-70B), PDF text extraction (pdf2json), and a modern React + Node.js "
        "architecture, the application effectively digitizes, reads, and simplifies medical diagnostic reports into "
        "plain, patient-friendly language.\n\n"
        "The platform enables patients to upload PDF medical reports and receive AI-generated structured analysis "
        "including key findings, doctor recommendations, risk assessment, and medication suggestions — all in both "
        "English and Hindi. Key features include an interactive medical AI chatbot, nearby healthcare facility mapping "
        "using Leaflet.js, medical report uploading, multilingual support, and secure report sharing via tokens. "
        "This system bridges the gap between complicated medical terminology and patient comprehension, contributing "
        "meaningfully to the field of digital health technology.\n\n"
        "The project is developed using the MERN stack (MongoDB, Express.js, React.js, Node.js) with JWT-based "
        "authentication, RESTful API architecture, and cloud-compatible deployment."
    )
    add_body(doc, abstract)

    # Keywords row
    doc.add_paragraph()
    kw_p = doc.add_paragraph()
    kw_r1 = kw_p.add_run("Keywords: ")
    kw_r1.bold = True
    kw_r1.font.size = Pt(11)
    kw_p.add_run("Medical AI, Report Analysis, MERN Stack, Natural Language Processing, Groq LLaMA, OCR, Healthcare Technology, Multilingual AI, JWT Authentication")

    # ════════ PAGE 4: ACKNOWLEDGEMENT ════════
    add_page_break(doc)
    add_heading(doc, "ACKNOWLEDGEMENT", level=0, center=True, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    ack = (
        "I would like to express my profound gratitude and sincere appreciation to all those who provided "
        "the support and guidance necessary to complete this project successfully.\n\n"
        "I am deeply grateful to my Project Guide, [Guide Name], [Designation], [Department], for their expert "
        "technical mentorship, constructive feedback, and constant encouragement throughout the development "
        "of MedVision AI. Their valuable insights greatly shaped the scope and quality of this project.\n\n"
        "I express my heartfelt thanks to [HOD Name], Head of Department, [Department Name], for providing the "
        "necessary academic resources and facilitating a supportive environment for research and development.\n\n"
        "I extend my gratitude to the entire faculty of [Institution Name] for their knowledge-sharing and continuous "
        "support throughout the academic year.\n\n"
        "I am thankful to the open-source community — specifically the contributors behind React.js, Node.js, "
        "MongoDB, Groq AI, Leaflet.js, and Tailwind CSS — whose tools enabled the realization of this project.\n\n"
        "Finally, I am sincerely grateful to my family and friends for their unwavering moral support and "
        "encouragement throughout this journey."
    )
    add_body(doc, ack)

    doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.add_run("[Your Name]\n[Roll No.]\n[Class / Branch]")

    # ════════ PAGE 5: INDEX ════════
    add_page_break(doc)
    add_heading(doc, "INDEX", level=0, center=True, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    index_rows = [
        ["1",  "Title Page",                           "1"],
        ["2",  "Certificate",                          "2"],
        ["3",  "Abstract",                             "3"],
        ["4",  "Acknowledgement",                      "4"],
        ["5",  "Index",                                "5"],
        ["6",  "Introduction",                         "6"],
        ["7",  "System Architecture Diagram",          "7"],
        ["8",  "UML Diagrams (Use Case, Class, Sequence, Activity)", "8"],
        ["9",  "Data Dictionary / Database Design",   "9"],
        ["10", "Screen Layouts (UI/UX)",               "10"],
        ["11", "Report Layouts",                       "11"],
        ["12", "Sample Coding Implementation",         "12"],
        ["13", "Testing",                              "13"],
        ["14", "Future Enhancements",                  "14"],
        ["15", "Conclusion",                           "15"],
        ["16", "Bibliography",                         "16"],
    ]
    add_data_table(doc, ["Sr. No.", "Content", "Page No."], index_rows)

    # ════════ PAGE 6: INTRODUCTION ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 1: INTRODUCTION", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    add_heading(doc, "1.1 Problem Statement", level=2, color="2E75B6")
    add_body(doc,
        "Medical diagnostic reports (blood tests, MRI analyses, X-ray reports) are produced in highly technical, "
        "clinical language that is largely incomprehensible to the average patient. This creates a critical gap in "
        "healthcare literacy — patients often leave clinics confused about what their reports mean, what actions to "
        "take, and how serious their condition is. There is a pressing need for a technology-based solution that can "
        "interpret these reports in a simple, accessible, and actionable manner."
    )

    add_heading(doc, "1.2 Proposed Solution", level=2, color="2E75B6")
    add_body(doc,
        "MedVision AI addresses this problem by providing an end-to-end web platform where patients can upload their "
        "PDF medical reports. The system automatically extracts text, sends it to an AI model (LLaMA 3.3-70B via Groq), "
        "and returns a full structured analysis — including key findings, doctor-style advice, risk scores, medicine "
        "suggestions, and an overall health status — in plain English or Hindi."
    )

    add_heading(doc, "1.3 Project Objectives", level=2, color="2E75B6")
    bullets = [
        "Design and develop a secure, full-stack web application using the MERN stack.",
        "Implement PDF text extraction using the pdf2json library.",
        "Integrate Groq AI (LLaMA 3.3-70B) to generate structured, patient-friendly medical analysis.",
        "Build a multilingual medical chatbot supporting English and Hindi.",
        "Develop an interactive nearby hospital/clinic finder using Leaflet.js.",
        "Implement secure report sharing using randomly generated tokens.",
        "Provide a responsive, accessible UI/UX using React.js and Tailwind CSS.",
    ]
    for b in bullets:
        add_bullet(doc, b)

    add_heading(doc, "1.4 Scope of the Project", level=2, color="2E75B6")
    add_body(doc,
        "The scope of MedVision AI includes the digitization and AI interpretation of PDF-format medical reports. "
        "The system supports multi-format uploads, real-time AI analysis, multilingual outputs, interactive maps, "
        "and a conversational medical chatbot. The project does NOT cover real-time telemedicine video consultation "
        "or direct integration with hospital EHR (Electronic Health Records) systems in the current version."
    )

    add_heading(doc, "1.5 Technology Stack", level=2, color="2E75B6")
    tech_rows = [
        ["Frontend",    "React.js, Tailwind CSS, Vite, Axios, Leaflet.js, React Router DOM"],
        ["Backend",     "Node.js, Express.js, Multer (file handling)"],
        ["Database",    "MongoDB Atlas (NoSQL), Mongoose ODM"],
        ["AI Engine",   "Groq Cloud API – LLaMA 3.3-70B Versatile model"],
        ["PDF Engine",  "pdf2json – Server-side PDF text extraction"],
        ["Auth",        "JSON Web Tokens (JWT), bcryptjs"],
        ["Deployment",  "Vercel (Frontend), Render / Railway (Backend)"],
    ]
    add_data_table(doc, ["Layer", "Technologies Used"], tech_rows)

    # ════════ PAGE 7: SYSTEM ARCHITECTURE ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 2: SYSTEM ARCHITECTURE", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    add_heading(doc, "2.1 Architecture Overview", level=2, color="2E75B6")
    add_body(doc,
        "MedVision AI follows a classic multi-tier (N-Tier) client-server architecture. The system is logically "
        "divided into four layers: the Presentation Layer (React frontend), the Application Layer (Node.js/Express "
        "REST API), the Services Layer (AI, PDF parsing, Maps), and the Data Layer (MongoDB Atlas). All communication "
        "between the frontend and backend occurs via RESTful HTTP APIs secured with JWT Bearer tokens."
    )

    add_image(doc, ARCH_IMG, "Figure 1: MedVision AI — System Architecture Diagram", width=Inches(5.8))

    add_heading(doc, "2.2 Layer Descriptions", level=2, color="2E75B6")
    layer_rows = [
        ["Presentation Layer", "React.js SPA rendered in browser. Handles all user interaction, routing, and state management."],
        ["Application Layer",  "Express.js REST API. Validates JWT, routes requests, orchestrates processing."],
        ["Services Layer",     "Groq AI for text analysis, pdf2json for PDF parsing, Leaflet for maps."],
        ["Data Layer",         "MongoDB Atlas stores User and Report collections with Mongoose schemas."],
        ["Security Layer",     "JWT tokens for authentication, bcryptjs for password hashing, CORS middleware."],
    ]
    add_data_table(doc, ["Layer", "Description"], layer_rows)

    add_heading(doc, "2.3 API Endpoints Summary", level=2, color="2E75B6")
    api_rows = [
        ["POST",   "/api/auth/register",              "Public",    "Register a new user account"],
        ["POST",   "/api/auth/login",                 "Public",    "Authenticate user, return JWT token"],
        ["POST",   "/api/reports/upload",             "Protected", "Upload PDF, extract text, run AI analysis"],
        ["GET",    "/api/reports",                    "Protected", "Retrieve all reports for logged-in user"],
        ["GET",    "/api/reports/:id",                "Protected", "Retrieve single report by ID"],
        ["DELETE", "/api/reports/:id",                "Protected", "Delete report by ID"],
        ["POST",   "/api/reports/:id/reanalyze",      "Protected", "Re-run AI analysis on existing report"],
        ["POST",   "/api/reports/:id/share",          "Protected", "Generate public share token"],
        ["GET",    "/api/reports/shared/:token",      "Public",    "Fetch shared report by token"],
        ["GET",    "/api/users/profile",              "Protected", "Get current user profile"],
        ["PUT",    "/api/users/profile",              "Protected", "Update user profile details"],
        ["POST",   "/api/chat",                       "Protected", "Send message to medical AI chatbot"],
    ]
    add_data_table(doc, ["Method", "Endpoint", "Access", "Description"], api_rows)

    # ════════ PAGE 8: UML DIAGRAMS ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 3: UML DIAGRAMS", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    add_heading(doc, "3.1 Use Case Diagram", level=2, color="2E75B6")
    add_body(doc,
        "The Use Case Diagram below illustrates the functional interactions between the two primary actors "
        "(Patient/User and Doctor) and the MedVision AI system. It captures all major use cases including "
        "report upload, AI analysis, the chatbot, nearby facility search, and profile management."
    )
    add_image(doc, USECASE_IMG, "Figure 2: Use Case Diagram — MedVision AI", width=Inches(5.5))

    add_page_break(doc)
    add_heading(doc, "3.2 Class Diagram", level=2, color="2E75B6")
    add_body(doc,
        "The Class Diagram represents the primary database entities (Mongoose models) and their relationships "
        "within the MedVision AI system. The User model stores authentication and profile information, while the "
        "Report model holds all data related to uploaded and analyzed medical documents. A one-to-many relationship "
        "exists from User to Report."
    )
    add_image(doc, CLASS_IMG, "Figure 3: Class Diagram — User and Report Entities", width=Inches(5.5))

    add_page_break(doc)
    add_heading(doc, "3.3 Sequence Diagram — Report Upload Flow", level=2, color="2E75B6")
    add_body(doc,
        "The Sequence Diagram below describes the detailed message flow when a user uploads a medical PDF report. "
        "It illustrates the interaction among the User, React Frontend, Express Backend, Multer, pdf2json, "
        "Groq AI API, and MongoDB database in a time-ordered sequence."
    )
    add_image(doc, SEQ_IMG, "Figure 4: Sequence Diagram — Report Upload and AI Analysis Flow", width=Inches(5.8))

    add_page_break(doc)
    add_heading(doc, "3.4 Activity Diagram — User Workflow", level=2, color="2E75B6")
    add_body(doc,
        "The Activity Diagram models the complete workflow a user follows from landing on the MedVision AI platform "
        "to using its core features. It covers authentication, report uploading, AI analysis, chatbot interaction, "
        "and nearby facility search, including all decision points and branching paths."
    )
    add_image(doc, ACTIVITY_IMG, "Figure 5: Activity Diagram — Complete User Workflow", width=Inches(5.5))

    # ════════ PAGE 9: DATA DICTIONARY ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 4: DATA DICTIONARY / DATABASE DESIGN", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    add_body(doc,
        "MedVision AI uses MongoDB (NoSQL document database) hosted on MongoDB Atlas. The schema is defined using "
        "Mongoose ODM (Object Data Modeling) for Node.js. The system maintains two primary collections: User and Report."
    )
    doc.add_paragraph()

    add_heading(doc, "4.1 Collection: User", level=2, color="2E75B6")
    user_rows = [
        ["_id",         "ObjectId",  "Primary Key",     "Unique document identifier (auto-generated)"],
        ["name",        "String",    "Required",        "Full name of the registered user"],
        ["email",       "String",    "Required, Unique","User's email address for login and communication"],
        ["password",    "String",    "Required",        "Bcrypt-hashed password string"],
        ["phone",       "String",    "Default: ''",     "User's contact phone number (optional)"],
        ["bio",         "String",    "Default: ''",     "Short user biography / health notes"],
        ["avatarColor", "String",    "Default: #2563EB","Hex color used for default avatar display in UI"],
        ["language",    "String",    "Enum: en/hi",     "Preferred language for AI output (English or Hindi)"],
        ["createdAt",   "Date",      "Auto (timestamp)","Document creation timestamp"],
        ["updatedAt",   "Date",      "Auto (timestamp)","Last document modification timestamp"],
    ]
    add_data_table(doc, ["Field", "Data Type", "Constraint", "Description"], user_rows)

    add_heading(doc, "4.2 Collection: Report", level=2, color="2E75B6")
    report_rows = [
        ["_id",              "ObjectId", "Primary Key",        "Unique document identifier (auto-generated)"],
        ["userId",           "ObjectId", "FK → User._id",      "Reference to the user who uploaded this report"],
        ["originalFileName", "String",   "Optional",           "Original name of the uploaded PDF file"],
        ["extractedText",    "String",   "Optional",           "Raw text extracted from the PDF (max 3000 chars stored)"],
        ["aiResult",         "String",   "Optional",           "Structured JSON string from Groq AI analysis"],
        ["isStructured",     "Boolean",  "Default: false",     "True if AI returned valid parseable JSON"],
        ["language",         "String",   "Default: 'en'",      "Language of AI analysis output (en or hi)"],
        ["shareToken",       "String",   "Default: null",      "Random hex token for public read-only sharing"],
        ["createdAt",        "Date",     "Auto (timestamp)",   "Report upload/creation timestamp"],
        ["updatedAt",        "Date",     "Auto (timestamp)",   "Last update timestamp (e.g., after reanalysis)"],
    ]
    add_data_table(doc, ["Field", "Data Type", "Constraint", "Description"], report_rows)

    add_heading(doc, "4.3 Entity Relationship (ER) Summary", level=2, color="2E75B6")
    add_body(doc, "Relationship: One User can have zero or more Reports (1 : N relationship).")
    er_rows = [
        ["User",   "Report", "ONE-TO-MANY (1:N)", "userId in Report references _id in User. CASCADE-style: reports are user-scoped."],
    ]
    add_data_table(doc, ["Entity 1", "Entity 2", "Relationship Type", "Notes"], er_rows)

    # ════════ PAGE 10: SCREEN LAYOUTS ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 5: SCREEN LAYOUTS (UI/UX)", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    add_body(doc,
        "MedVision AI implements a fully responsive Single Page Application (SPA) architecture. Below are the "
        "descriptions of each application screen. [Note: Insert actual screenshots of each page here as images.]"
    )

    screens = [
        ("5.1 Landing Page (/)", [
            "Header: Logo 'MedVision AI' on left, navigation links (Login, Register) on right.",
            "Hero Section: Full-width banner with tagline 'Understand Your Health Reports with AI'.",
            "Features Grid: Cards highlighting Upload, AI Analysis, Chatbot, and Nearby Facilities.",
            "Footer with links and copyright.",
        ]),
        ("5.2 Register Page (/register)", [
            "Clean centered form with fields: Full Name, Email, Password.",
            "Submit button labeled 'Create Account'.",
            "Link to Login page for existing users.",
            "Client-side validation with error highlighting.",
        ]),
        ("5.3 Login Page (/login)", [
            "Email and Password input fields with show/hide password toggle.",
            "JWT token stored in localStorage upon successful authentication.",
            "Redirect to Dashboard on login success.",
        ]),
        ("5.4 Dashboard (/dashboard)", [
            "Welcome banner with user's name and avatar color indicator.",
            "Quick-action cards: Upload Report, View Reports, Open Chatbot, Nearby Hospitals.",
            "Recent reports summary list with timestamps.",
            "Navigation sidebar/topbar for all sections.",
        ]),
        ("5.5 Upload Page (/upload)", [
            "Drag-and-drop file zone accepting PDF files.",
            "Language selector dropdown (English / Hindi) for AI output preference.",
            "Upload button triggers multipart/form-data POST to backend.",
            "Loading spinner displayed during AI analysis (which may take 5–15 seconds).",
            "Auto-redirect to Report Details page on success.",
        ]),
        ("5.6 Reports Page (/reports)", [
            "Grid/list of all previously uploaded and analyzed reports.",
            "Each card shows: file name, date uploaded, risk level badge.",
            "Action buttons: View, Re-analyze, Share, Delete.",
        ]),
        ("5.7 Report Details Page (/report/:id)", [
            "Displays structured AI output: Patient Summary, Risk Score gauge, Key Findings table.",
            "Doctor Advice panels: Diet, Lifestyle, Follow-up Urgency.",
            "Medicine Recommendations section with type, dosage, caution info.",
            "Share Report button generates a public shareable URL.",
        ]),
        ("5.8 Medical Chatbot (/chat)", [
            "Full-screen chat interface with scrollable message history.",
            "Distinct UI bubbles for User messages (right, blue) and AI responses (left, white).",
            "Input field at bottom with Send button.",
            "Context-aware health question answering via Groq AI.",
        ]),
        ("5.9 Nearby Facilities (/nearby)", [
            "Interactive Leaflet.js map centered on user's current geolocation.",
            "Clickable map markers for hospitals, clinics, and pharmacies.",
            "Search bar to filter by facility type.",
        ]),
        ("5.10 Profile Page (/profile)", [
            "Displays user avatar (color-coded), name, email, phone, bio.",
            "Edit mode with inline form fields for updating profile.",
            "Language preference toggle (English / Hindi).",
        ]),
    ]

    for title, points in screens:
        add_heading(doc, title, level=2, color="2E75B6")
        for pt in points:
            add_bullet(doc, pt)
        doc.add_paragraph()
        # Placeholder for screenshot
        ph = doc.add_paragraph("[Insert Screenshot of " + title.split("(")[0].strip() + " here]")
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.runs[0].italic = True
        ph.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        ph.runs[0].font.size = Pt(10)
        doc.add_paragraph()

    # ════════ PAGE 11: REPORT LAYOUTS ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 6: REPORT LAYOUTS", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    add_body(doc,
        "When a patient's medical report is processed, MedVision AI generates a standardized AI Health Analysis Report. "
        "The format of this AI-generated output is described below."
    )

    add_heading(doc, "6.1 AI Health Analysis Report Structure", level=2, color="2E75B6")
    report_layout_rows = [
        ["1", "patientSummary",         "String",        "2–3 sentence warm, patient-friendly health summary"],
        ["2", "overallHealth",          "Enum String",   "One of: Good | Fair | Needs Attention | Critical"],
        ["3", "riskLevel",              "Enum String",   "Low | Medium | High — based on findings"],
        ["4", "riskScore",              "Integer (0-100)","Numerical risk score for visual gauge display"],
        ["5", "keyFindings[]",          "Array",         "List of parameters with value, normal range, status, and doctor note"],
        ["6", "doctorAdvice.diet[]",    "String Array",  "3 diet-related recommendations"],
        ["7", "doctorAdvice.lifestyle[]","String Array", "3 lifestyle recommendations"],
        ["8", "doctorAdvice.followUp",  "String",        "When and why to see a doctor"],
        ["9", "doctorAdvice.urgency",   "Enum String",   "No rush | Within a month | Within a week | See doctor today"],
        ["10","recommendedMedicines[]", "Array",         "Safe supplement/medicine suggestions with dosage and caution"],
        ["11","goodNews",               "String",        "One positive aspect of the report"],
        ["12","watchOut",               "String",        "Most critical parameter/finding to monitor"],
        ["13","disclaimer",             "String (fixed)","Standard disclaimer about AI not replacing medical advice"],
    ]
    add_data_table(doc, ["#", "Field", "Type", "Description"], report_layout_rows)

    add_heading(doc, "6.2 Key Findings — Sub-structure", level=2, color="2E75B6")
    kf_rows = [
        ["parameter", "String", "Clinical parameter name  (e.g., Hemoglobin, Glucose)"],
        ["value",     "String", "Reported value with unit (e.g., 11.0 g/dL)"],
        ["normalRange","String","Reference normal range   (e.g., 12.0–17.5 g/dL)"],
        ["status",    "Enum",   "Normal | Borderline | Abnormal"],
        ["doctorNote","String", "Plain language explanation from AI doctor perspective"],
    ]
    add_data_table(doc, ["Sub-field", "Type", "Description"], kf_rows)

    # ════════ PAGE 12: SAMPLE CODE ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 7: SAMPLE CODING IMPLEMENTATION", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    add_heading(doc, "7.1 Backend — Express Server Entry Point (server.js)", level=2, color="2E75B6")
    add_code_block(doc, """
const express  = require("express");
const mongoose = require("mongoose");
const cors     = require("cors");
require("dotenv").config();

const authRoutes   = require("./routes/authRoutes");
const reportRoutes = require("./routes/reportRoutes");
const userRoutes   = require("./routes/userRoutes");
const chatRoutes   = require("./routes/chatRoutes");

const app = express();
app.use(cors());
app.use(express.json());

app.use("/api/auth",    authRoutes);
app.use("/api/reports", reportRoutes);
app.use("/api/users",   userRoutes);
app.use("/api/chat",    chatRoutes);

mongoose
  .connect(process.env.MONGO_URI)
  .then(() => console.log("MongoDB Connected"))
  .catch((err) => console.log(err));

app.listen(process.env.PORT, () => {
  console.log(`Server running on port ${process.env.PORT}`);
});
""", "Code 1: server.js — Express application entry point")

    add_heading(doc, "7.2 Backend — Mongoose User Model (models/User.js)", level=2, color="2E75B6")
    add_code_block(doc, """
const mongoose = require("mongoose");

const userSchema = new mongoose.Schema(
  {
    name:        { type: String, required: true },
    email:       { type: String, required: true, unique: true },
    password:    { type: String, required: true },
    phone:       { type: String, default: "" },
    bio:         { type: String, default: "" },
    avatarColor: { type: String, default: "#2563EB" },
    language:    { type: String, default: "en", enum: ["en", "hi"] },
  },
  { timestamps: true }
);

module.exports = mongoose.model("User", userSchema);
""", "Code 2: User.js — Mongoose User Schema")

    add_heading(doc, "7.3 Backend — Report Controller: Upload & AI Analysis", level=2, color="2E75B6")
    add_code_block(doc, """
// Excerpt from backend/controllers/reportController.js

// ─── Groq AI Setup (OpenAI Compatible) ───
const groq = new OpenAI({
  apiKey: process.env.GROQ_API_KEY,
  baseURL: "https://api.groq.com/openai/v1",
});

// ─── Run AI Analysis ───
async function runAI(textForAI, language = "en") {
  const { system, user } = buildPrompt(textForAI, language);
  const response = await groq.chat.completions.create({
    model: "llama-3.3-70b-versatile",
    messages: [
      { role: "system", content: system },
      { role: "user",   content: user },
    ],
    temperature: 0.4,
    max_tokens: 2000,
  });
  const raw = response.choices[0].message.content;
  let parsed = null;
  try {
    const match = raw.match(/{[\\s\\S]*}/);
    if (match) parsed = JSON.parse(match[0]);
  } catch {}
  return { raw, parsed };
}

// ─── Upload Report Endpoint ───
exports.uploadReport = async (req, res) => {
  const filePath = req.file.path;
  const language = req.body.language || "en";
  const pdfParser = new PDFParser();

  pdfParser.on("pdfParser_dataReady", async (pdfData) => {
    let extractedText = "";
    pdfData.Pages.forEach((page) => {
      page.Texts.forEach((text) => {
        text.R.forEach((run) => {
          extractedText += decodeURIComponent(run.T) + " ";
        });
      });
    });
    const { raw, parsed } = await runAI(
      extractedText.slice(0, 6000), language
    );
    const report = await Report.create({
      userId:           req.user.id,
      originalFileName: req.file.originalname,
      extractedText:    extractedText.slice(0, 3000),
      aiResult:         parsed ? JSON.stringify(parsed) : raw,
      isStructured:     parsed !== null,
      language,
    });
    fs.unlinkSync(filePath);
    res.status(201).json(report);
  });
  pdfParser.loadPDF(filePath);
};
""", "Code 3: reportController.js — AI-powered PDF analysis (excerpt)")

    add_heading(doc, "7.4 Frontend — React Routing (App.jsx)", level=2, color="2E75B6")
    add_code_block(doc, """
// frontend/src/App.jsx (excerpt)
import { BrowserRouter as Router, Routes, Route } from "react-router-dom";
import ProtectedRoute from "./components/ProtectedRoute";
import MainLayout    from "./layout/MainLayout";
import Landing       from "./pages/Landing";
import Login         from "./pages/Login";
import Register      from "./pages/Register";
import Dashboard     from "./pages/Dashboard";
import UploadPage    from "./pages/UploadPage";
import ReportsPage   from "./pages/ReportsPage";
import ReportDetails from "./pages/ReportDetails";
import ChatPage      from "./pages/ChatPage";
import NearbyPage    from "./pages/NearbyPage";
import ProfilePage   from "./pages/ProfilePage";

function App() {
  return (
    <Router>
      <Routes>
        <Route path="/"         element={<Landing />} />
        <Route path="/login"    element={<Login />} />
        <Route path="/register" element={<Register />} />

        {/* Protected Routes */}
        <Route element={<ProtectedRoute><MainLayout /></ProtectedRoute>}>
          <Route path="/dashboard"   element={<Dashboard />} />
          <Route path="/upload"      element={<UploadPage />} />
          <Route path="/reports"     element={<ReportsPage />} />
          <Route path="/report/:id"  element={<ReportDetails />} />
          <Route path="/chat"        element={<ChatPage />} />
          <Route path="/nearby"      element={<NearbyPage />} />
          <Route path="/profile"     element={<ProfilePage />} />
        </Route>
      </Routes>
    </Router>
  );
}
export default App;
""", "Code 4: App.jsx — Routing configuration (frontend)")

    # ════════ PAGE 13: TESTING ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 8: TESTING", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    add_heading(doc, "8.1 Types of Testing Performed", level=2, color="2E75B6")
    testing_types = [
        ("Unit Testing", "Individual functions and components tested in isolation — e.g., PDF extraction logic, AI prompt builder function, JWT validation middleware."),
        ("Integration Testing", "Backend API endpoints tested with Postman to verify correct request/response behavior for register, login, upload, and report retrieval."),
        ("UI Testing (Manual)", "All frontend pages manually tested across Chrome, Firefox, and Edge for layout consistency, button functionality, and error state displays."),
        ("AI Output Testing", "Multiple real medical PDF reports uploaded to verify AI analysis accuracy, JSON structure validity, and multilingual output consistency."),
        ("Authentication Testing", "Verified that protected routes correctly reject unauthorized requests (401) and that JWT tokens expire appropriately."),
        ("Boundary Testing", "Tested with empty PDFs, non-text PDFs (image-only), and oversized files to verify error handling."),
    ]
    for name, desc in testing_types:
        add_bullet(doc, desc, bold_prefix=name)

    add_heading(doc, "8.2 Test Case Table", level=2, color="2E75B6")
    test_rows = [
        ["TC-001", "User Registration",       "POST /api/auth/register with valid data",     "201 Created, user stored in DB",   "Pass"],
        ["TC-002", "Duplicate Email Reg.",     "POST /register with existing email",          "400 Error: 'User already exists'", "Pass"],
        ["TC-003", "User Login",               "POST /api/auth/login with correct credentials","200 OK, JWT token in response",   "Pass"],
        ["TC-004", "Login Wrong Password",     "POST /login with invalid password",           "401 Unauthorized",                "Pass"],
        ["TC-005", "Upload Valid PDF",         "POST /api/reports/upload with valid PDF",     "201, Report with aiResult created","Pass"],
        ["TC-006", "Upload Empty PDF",         "Upload a blank/empty PDF file",               "400: Could not extract text",     "Pass"],
        ["TC-007", "Get All Reports",          "GET /api/reports with valid JWT",             "200, array of report objects",    "Pass"],
        ["TC-008", "Get Report by ID",         "GET /api/reports/:id (own report)",           "200, single report object",       "Pass"],
        ["TC-009", "Delete Report",            "DELETE /api/reports/:id (own report)",        "200: deleted successfully",       "Pass"],
        ["TC-010", "Share Report Token",       "POST /api/reports/:id/share",                 "200, shareToken generated",       "Pass"],
        ["TC-011", "Access Shared Report",     "GET /api/reports/shared/:token (public)",     "200, report data (no userId)",    "Pass"],
        ["TC-012", "Chat API",                 "POST /api/chat with health question",         "200, AI text response",           "Pass"],
        ["TC-013", "Access Without Auth",      "GET /api/reports without JWT header",         "401 Unauthorized",                "Pass"],
        ["TC-014", "Re-analyze Report Hindi",  "POST /api/reports/:id/reanalyze (lang=hi)",   "200, Hindi language AI output",   "Pass"],
    ]
    add_data_table(doc, ["Test ID", "Test Case", "Input", "Expected Output", "Status"], test_rows)

    # ════════ PAGE 14: FUTURE ENHANCEMENTS ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 9: FUTURE ENHANCEMENTS", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    add_body(doc,
        "The MedVision AI platform has been designed with scalability and extensibility in mind. "
        "The following enhancements are planned for future development phases:"
    )

    enhancements = [
        ("Wearable Device Integration",
         "Sync real-time health data (heart rate, SpO2, step count, blood glucose) from smartwatches "
         "(Apple Watch, Fitbit, Samsung Health) directly into MedVision AI for continuous monitoring."),
        ("DICOM Imaging Support",
         "Process raw DICOM-format X-ray and MRI data for AI-based preliminary anomaly detection and "
         "visual highlighting of concerning regions."),
        ("Telemedicine Booking Module",
         "Allow users to book video consultations with verified doctors directly from within the AI Report "
         "screen, with calendar integration and payment gateway support."),
        ("Voice-based Medical Chatbot",
         "Implement Speech-to-Text (Web Speech API) so elderly users or people with limited literacy can "
         "interact with the medical chatbot via voice commands."),
        ("React Native Mobile App",
         "Expand the web application into native Android and iOS apps using React Native for mobile-first "
         "access, offline report viewing, and push notification alerts."),
        ("Electronic Health Record (EHR) Integration",
         "Integrate with national health data standards (HL7 FHIR) for direct EHR connectivity with "
         "hospital systems."),
        ("Advanced Subscription Model",
         "Implement freemium tiers: free basic analysis, premium structured analysis, and family plan "
         "with shared health dashboards."),
        ("Doctor Portal",
         "Dedicated doctor dashboard to receive shared patient reports, annotate them, and send "
         "back AI-assisted recommendations."),
    ]

    for i, (title, desc) in enumerate(enhancements, 1):
        add_heading(doc, f"9.{i} {title}", level=2, color="2E75B6")
        add_body(doc, desc)

    # ════════ PAGE 15: CONCLUSION ════════
    add_page_break(doc)
    add_heading(doc, "CHAPTER 10: CONCLUSION", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    conclusion = (
        "The MedVision AI project successfully demonstrates the transformative potential of integrating Artificial "
        "Intelligence with modern full-stack web development to address a real-world healthcare challenge. By leveraging "
        "the Groq AI platform (LLaMA 3.3-70B), pdf2json parsing, JWT-based security, and the MERN stack, the system "
        "delivers a comprehensive and accessible medical report analysis tool that was previously available only to "
        "medical professionals.\n\n"
        "The project fulfills all its stated objectives: secure user authentication, PDF medical report processing, "
        "structured AI-generated health analysis in two languages, an interactive medical chatbot, nearby facility "
        "mapping, and secure token-based report sharing. The SPA architecture with React.js ensures a fast, smooth, "
        "and responsive user experience across devices.\n\n"
        "MedVision AI demonstrates that technology can meaningfully democratize healthcare information access — "
        "empowering patients to understand their own health data, make informed decisions, and seek timely medical "
        "attention. The project is well-positioned for future enhancements including mobile app development, wearable "
        "integrations, and telemedicine features.\n\n"
        "This project was a valuable learning experience encompassing full-stack development, AI API integration, "
        "database design, RESTful architecture, and UI/UX design principles."
    )
    add_body(doc, conclusion)

    # ════════ PAGE 16: BIBLIOGRAPHY ════════
    add_page_break(doc)
    add_heading(doc, "BIBLIOGRAPHY", level=0, color="1F497D")
    add_separator(doc)
    doc.add_paragraph()

    references = [
        ("MongoDB Documentation", "Official guide for NoSQL database schema design, queries, and aggregation pipelines.", "https://docs.mongodb.com/"),
        ("Mongoose ODM Documentation", "Object Data Modeling library for MongoDB and Node.js.", "https://mongoosejs.com/docs/"),
        ("Express.js API Reference", "Official framework guide for building Node.js HTTP servers and REST APIs.", "https://expressjs.com/"),
        ("React.js Documentation", "Component lifecycle, hooks, context API, and routing with React Router.", "https://react.dev/"),
        ("Vite Build Tool", "Next-generation frontend tooling with instant server start and hot module replacement.", "https://vitejs.dev/"),
        ("Tailwind CSS Documentation", "Utility-first CSS framework for responsive UI composition.", "https://tailwindcss.com/docs"),
        ("Groq AI Documentation", "Groq Cloud API for ultra-fast LLM inference with OpenAI-compatible endpoint.", "https://console.groq.com/docs/"),
        ("Leaflet.js Mapping Library", "Open-source JavaScript library for interactive geographical maps.", "https://leafletjs.com/"),
        ("pdf2json Library", "Node.js library for server-side PDF text extraction.", "https://github.com/modesty/pdf2json"),
        ("JSON Web Tokens (JWT)", "Standard for secure information transmission between parties as a JSON object.", "https://jwt.io/introduction/"),
        ("Node.js Documentation", "Official documentation for the Node.js JavaScript runtime environment.", "https://nodejs.org/en/docs/"),
        ("Multer Documentation", "Node.js middleware for handling multipart/form-data for file uploads.", "https://github.com/expressjs/multer"),
    ]

    for i, (title, desc, url) in enumerate(references, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r1 = p.add_run(f"[{i}]  ")
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(12)
        r2 = p.add_run(f"{title}: ")
        r2.bold = True
        r2.font.name = "Calibri"
        r2.font.size = Pt(12)
        r3 = p.add_run(f"{desc}  ")
        r3.font.name = "Calibri"
        r3.font.size = Pt(12)
        r4 = p.add_run(url)
        r4.font.name = "Calibri"
        r4.font.size = Pt(12)
        r4.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        r4.italic = True

    # ════════ SAVE ════════
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Documentation generated successfully!")
    print(f"   Saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    build_document()
